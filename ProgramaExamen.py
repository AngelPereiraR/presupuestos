import docx
import docx2pdf

menuPrincipal = "-"
costeMensualHoraBeneficio = 0

while(menuPrincipal != "SALIR"):
    menuPrincipal = input("¿Deseas introducir el coste de una hora de la mano de obra manualmente? (S/N): ")

    if(menuPrincipal == "S" or menuPrincipal == "s"):
        costeMensualHoraBeneficio = float(input("Introduce el valor del coste de la hora: "))
        menuPrincipal = "SALIR"
    elif(menuPrincipal == "N" or menuPrincipal == "n"):
        opcionManoObra = "0"
        costeTotalMensual = 0
        
        while(opcionManoObra != "3"):
            print("---- Menú Coste/hora de la mano de obra ----")
            print("1. Añadir coste.")
            print("2. Calcular coste por hora con beneficio.")
            print("3. Salir.")
            opcionManoObra = input("Escoge la acción a realizar: ")

            if(opcionManoObra == "1"):
                opcionTipo = "0"
                while(opcionTipo != "1" and opcionTipo != "2" and opcionTipo != "3" and opcionTipo != "4" and opcionTipo != "5" and opcionTipo != "6" and opcionTipo != "7" and opcionTipo != "8"):
                    print("---- Menú Tipos de coste ----")
                    print("1. Añadir coste normal (Coste anual, semestral, cuatrimestral, trimestral, bimensual, mensual, bisemanal, semanal o diario directo. Ej: 25€/semana).")
                    print("2. Añadir coste de sueldo.")
                    print("3. Añadir coste de seguridad social (Autónomo o trabajador).")
                    print("4. Añadir coste de amortización (Coste con Valor Económico, Valor Residual y/o años de Vida Útil. Ej: VE:3000€, VU:5 años).")
                    print("5. Añadir coste por número de días en específico (Ej: 30€/22 días).")
                    print("6. Añadir coste por número de semanas en específico (Ej: 120€/3 semanas).")
                    print("7. Añadir coste por número de meses en específico (Ej: 500€/7 meses).")
                    print("8. Añadir coste por número de años en específico (Ej: 3000€/2 años).")
                    opcionTipo = input ("Escoge la acción a realizar: ")

                    valorTransaccion = float(input("Introduce el valor del coste en euros: "))

                    valorMensual = 0

                    if(opcionTipo == "1"):
                        opcionTipoNormal = "0"
                        while(opcionTipoNormal != "1" and opcionTipoNormal != "2" and opcionTipoNormal != "3" and opcionTipoNormal != "4" and opcionTipoNormal != "5" and opcionTipoNormal != "6" and opcionTipoNormal != "7" and opcionTipoNormal != "8" and opcionTipoNormal != "9"):
                            print("---- Menú Tipos de coste normal ----")
                            print("1. Añadir coste anual.")
                            print("2. Añadir coste semestral.")
                            print("3. Añadir coste cuatrimestral.")
                            print("4. Añadir coste trimensual.")
                            print("5. Añadir coste bimensual.")
                            print("6. Añadir coste mensual.")
                            print("7. Añadir coste bisemanal.")
                            print("8. Añadir coste semanal.")
                            print("9. Añadir coste diario.")
                            opcionTipoNormal = input("Escoge la acción a realizar: ")

                            if(opcionTipoNormal == "1"):
                                valorMensual = valorTransaccion / 12
                            elif(opcionTipoNormal == "2"):
                                valorMensual = valorTransaccion / 6
                            elif(opcionTipoNormal == "3"):
                                valorMensual = valorTransaccion / 4
                            elif(opcionTipoNormal == "4"):
                                valorMensual = valorTransaccion / 3
                            elif(opcionTipoNormal == "5"):
                                valorMensual = valorTransaccion / 2
                            elif(opcionTipoNormal == "6"):
                                valorMensual = valorTransaccion
                            elif(opcionTipoNormal == "7"):
                                valorMensual = valorTransaccion * 2
                            elif(opcionTipoNormal == "8"):
                                valorMensual = valorTransaccion * 4
                            elif(opcionTipoNormal == "9"):
                                valorMensual = valorTransaccion * 30
                            else:
                                print("Lo siento, el valor introducido no corresponde a ninguna acción. Volviendo a mostrar el menú...")

                            clausula = ""
                            dineroExtra = 0

                            while(clausula != "S" and clausula != "s" and clausula != "N" and clausula != "n"):
                                clausula = input("Hay alguna cláusula por número de nóminas? (S/N): ")
                                if(clausula == "S" or clausula == "s"):
                                    trabajadores = input("¿Cuántos trabajadores están trabajando en el encargo?: ")
                                    dineroNomina = input("¿Cuánto dinero hay que aumentar por nómina?: ")
                                    dineroExtra = int(trabajadores) * float(dineroNomina)
                                elif(clausula == "N" or clausula == "n"):
                                    dineroExtra = 0
                                else:
                                    print("No es una opción válida, pruebe de nuevo.")
                                
                            valorMensual += dineroExtra
                    elif(opcionTipo == "2"):
                        valorMensual = valorTransaccion
                    elif(opcionTipo == "3"):
                        opcionSeguridad = "0"

                        while(opcionSeguridad != "1" and opcionSeguridad != "2"):
                            print("---- Menú Coste Seguridad Social ----")
                            print("1. Autónomo.")
                            print("2. Trabajador.")
                            opcionSeguridad = input("Escoge la acción a realizar: ")

                            if(opcionSeguridad == "1"):
                                valorMensual = float(valorTransaccion)
                            elif(opcionSeguridad == "2"):
                                horasTrabajadorNormal = input("¿Cuántas horas debe trabajar diariamente normalmente un trabajador?: ")
                                horasTrabajador = input("¿Cuántas horas trabaja diariamente el trabajador?: ")

                                valorMensual = valorTransaccion * int(horasTrabajador) / int(horasTrabajadorNormal)
                            else:
                                print("Lo siento, el valor introducido no corresponde a ninguna acción. Volviendo a mostrar el menú...")
                    elif(opcionTipo == "4"):
                        vr = ""

                        while(vr != "S" and vr != "s" and vr != "N" and vr != "n"):
                            vr = input("¿El coste tiene VR? (S/N): ")
                            
                            if(vr == "S" or vr == "s"):
                                valorVr = float(input("Introduce el valor de VR: "))
                            elif(vr == "N" or vr == "n"):
                                valorVr = 0
                            else:
                                print("No es una opción válida, pruebe de nuevo.")

                        vu = int(input("¿Cuántos años de pago tiene?: "))
                            
                        valorMensual = (valorTransaccion - valorVr) / vu / 12
                    elif(opcionTipo == "5"):
                        dias = input("Introduce el número de días: ")
                        valorMensual = valorTransaccion / int(dias) / 30
                    elif(opcionTipo == "6"):
                        semanas = input("Introduce el número de semanas: ")
                        valorMensual = valorTransaccion * int(semanas) / 4
                    elif(opcionTipo == "7"):
                        meses = input("Introduce el número de meses: ")
                        valorMensual = valorTransaccion / int(meses)
                    elif(opcionTipo == "8"):
                        anyos = input("Introduce el número de años: ")
                        valorMensual = valorTransaccion / int(anyos) / 12
                    else:
                        print("Lo siento, el valor introducido no corresponde a ninguna acción. Volviendo a mostrar el menú...")
                    
                    costeTotalMensual += valorMensual
            elif(opcionManoObra == "2" and costeTotalMensual != "0"):
                beneficioPorHora = input("¿Cuánto debe ser el beneficio por hora?: ")
                diasMes = input("¿Cuántos días se trabajan al mes?: ")
                numTrabajadores = input("¿Cuántos trabajadores (incluído el autónomo) trabajan en el encargo?: ")
                numTrabajadoresBucle = int(numTrabajadores)
                horasTrabajadoresDia = 0
                for i in range(numTrabajadoresBucle):
                    if(i == 0):
                        horasTrabajadoresDia += int(input("¿Cuántas horas trabaja el autónomo diariamente?: "))
                    else:
                        horasTrabajadoresDia += int(input("¿Cuántas horas trabaja el trabajador diariamente?: "))
                descanso = input("¿Cuántos minutos de descanso tiene cada trabajador?: ")
                horasTotalesDia = horasTrabajadoresDia - float(int(numTrabajadores) * int(descanso) / 60)
                horasTotalesMes = horasTotalesDia * int(diasMes)
                costeHoraMensual = float(costeTotalMensual) / horasTotalesMes
                costeMensualHoraBeneficio = round(costeHoraMensual + float(beneficioPorHora), 2)

                print("El coste por hora de mano de obra es:", costeMensualHoraBeneficio)
            elif(opcionManoObra == "2" and costeTotalMensual == "0"):
                print("Lo siento, el valor de coste total mensual debe ser mayor que 0 para poder realizar el cálculo. Volviendo a mostrar el menú...")
            elif(opcionManoObra == "3" and costeMensualHoraBeneficio != "0"):
                menuPrincipal = "SALIR"
                print("Saliendo del menú...")
            elif(opcionManoObra == "3" and costeMensualHoraBeneficio == "0"):
                print("Lo siento, no se puede salir del menú hasta que no se calcule el coste por hora. Volviendo a mostrar el menú...")
            else:
                print("Lo siento, el valor introducido no corresponde a ninguna acción. Volviendo a mostrar el menú...")

    else:
        print("No es una opción válida, pruebe de nuevo.")

    if(menuPrincipal == "SALIR"):
        #Presupuesto

        datosVendedor = {}
        datosVendedor["nombre"] = input("Introduce el nombre del vendedor: ")
        cifOrNif = ""
        while(cifOrNif != "CIF" and cifOrNif != "NIF"):
            cifOrNif = input("¿Es CIF o NIF?: ")
            if(cifOrNif == "CIF"):
                datosVendedor["cif"] = input("Introduce el CIF: ")
            elif(cifOrNif == "NIF"):
                datosVendedor["nif"] = input("Introduce el NIF: ")
            else:
                print("La opción no es válida, pruebe de nuevo.")
        datosVendedor["calle"] = input("Introduce el nombre de la calle: ")
        datosVendedor["codigoPostal"] = input("Introduce el código postal y la localidad: ")

        datosCliente = {}
        datosCliente["nombre"] = input("Introduce el nombre del cliente: ")
        cifOrNif = ""
        while(cifOrNif != "CIF" and cifOrNif != "NIF"):
            cifOrNif = input("¿Es CIF o NIF?: ")
            if(cifOrNif == "CIF"):
                datosCliente["cif"] = input("Introduce el CIF: ")
            elif(cifOrNif == "NIF"):
                datosCliente["nif"] = input("Introduce el NIF: ")
            else:
                print("La opción no es válida, pruebe de nuevo.")
        datosCliente["calle"] = input("Introduce el nombre de la calle: ")
        datosCliente["codigoPostal"] = input("Introduce el código postal y la localidad: ")

        datosPresupuesto = {}
        datosPresupuesto["numero"] = input("Introduce el número del presupuesto: ")
        datosPresupuesto["fecha"] = input("Introduce la fecha del presupuesto: ")
        datosPresupuesto["descripcion"] = input("Introduce la descripción del presupuesto: ")
        datosPresupuesto["observaciones"] = input("Introduce las observaciones del presupuesto: ")
        datosPresupuesto["clausulas"] = input("Introduce las cláusulas del presupuesto: ")

        anadirArticuloPresupuesto = ""
        articulosPresupuesto = []
        while(anadirArticuloPresupuesto != "2"):
            print("---- Menú Artículos Presupuesto ----")
            print("1. Añadir.")
            print("2. Salir.")
            anadirArticuloPresupuesto = input("Escoge la acción a realizar: ")

            if(anadirArticuloPresupuesto == "1"):
                ref = input("Introduce la referencia del artículo: ")
                nombre = input("Introduce el nombre del artículo: ")
                cantidad = float(input("Introduce la cantidad del artículo en números: "))
                precio = float(input("Introduce el precio del artículo en números: "))
                descuento = float(input("Introduce el valor del descuento del artículo (de 0 a 100): "))
                importe = precio * cantidad - (precio * cantidad * descuento / 100)
                articulosPresupuesto.append([ref, nombre, cantidad, precio, descuento, round(importe, 2)])
            elif(anadirArticuloPresupuesto == "2"):
                print("Saliendo del menú...")
            else:
                print("Lo siento, el valor introducido no corresponde a ninguna acción. Volviendo a mostrar el menú...")
        
        horasManoObra = int(input("Introduce el número de horas que trabajará la mano de obra: "))
        importeManoObra = round(horasManoObra * costeMensualHoraBeneficio, 2)

        baseImponible = 0
        for articulo in articulosPresupuesto:
            baseImponible += articulo[5]
        baseImponible += importeManoObra
        
        iva = float(input("Introduce el valor del IVA: "))
        importeIva = round(baseImponible * iva / 100, 2)

        totalPresupuesto = round(baseImponible - importeIva, 2)

        print("El total del presupuesto es:", totalPresupuesto)

        presupuesto = docx.Document()
        table = presupuesto.add_table(rows=2, cols=8)
        table.style = 'Table Grid'
        if(datosVendedor.get("cif") != None):
            table.rows[0].cells[0].text = "\nDatos del vendedor:\n" + datosVendedor.get("nombre") + "\nCIF: " + datosVendedor.get("cif") + "\n" + datosVendedor.get("calle") + "\nCP " + datosVendedor.get("codigoPostal") + "\n"
        else:
            table.rows[0].cells[0].text = "\nDatos del vendedor:\n" + datosVendedor.get("nombre") + "\nNIF: " + datosVendedor.get("nif") + "\n" + datosVendedor.get("calle") + "\nCP " + datosVendedor.get("codigoPostal") + "\n"
        
        table.rows[0].cells[4].text = "\nPresupuesto Nº " + datosPresupuesto.get("numero") + "\nRef.\nFecha " + datosPresupuesto.get("fecha") + "\n"
        table.rows[1].cells[0].text = "\nDescripción:\n" + datosPresupuesto.get("descripcion") + "\n"
        if(datosCliente.get("cif") != None):
            table.rows[1].cells[4].text = "\nDatos del cliente:\n" + datosCliente.get("nombre") + "\nCIF: " + datosCliente.get("cif") + "\n" + datosCliente.get("calle") + "\nCP " + datosCliente.get("codigoPostal") + "\n"
        else:
            table.rows[1].cells[4].text = "\nDatos del cliente:\n" + datosCliente.get("nombre") + "\nNIF: " + datosCliente.get("nif") + "\n" + datosCliente.get("calle") + "\nCP " + datosCliente.get("codigoPostal") + "\n"
        table.rows[0].cells[0].merge(table.rows[0].cells[1])
        table.rows[0].cells[0].merge(table.rows[0].cells[2])
        table.rows[0].cells[0].merge(table.rows[0].cells[3])
        table.rows[0].cells[4].merge(table.rows[0].cells[5])
        table.rows[0].cells[4].merge(table.rows[0].cells[6])
        table.rows[0].cells[4].merge(table.rows[0].cells[7])
        table.rows[1].cells[0].merge(table.rows[1].cells[1])
        table.rows[1].cells[0].merge(table.rows[1].cells[2])
        table.rows[1].cells[0].merge(table.rows[1].cells[3])
        table.rows[1].cells[4].merge(table.rows[1].cells[5])
        table.rows[1].cells[4].merge(table.rows[1].cells[6])
        table.rows[1].cells[4].merge(table.rows[1].cells[7])

        table.add_row()
        table.rows[2].cells[0].text = "Ref"
        table.rows[2].cells[1].text = "Artículos"
        table.rows[2].cells[1].merge(table.rows[2].cells[2])
        table.rows[2].cells[1].merge(table.rows[2].cells[3])
        table.rows[2].cells[4].text = "Cantidad"
        table.rows[2].cells[5].text = "Precio"
        table.rows[2].cells[6].text = "Dto %"
        table.rows[2].cells[7].text = "Importe"

        table.add_row()
        ref = "\n"
        articulos = "\n"
        cantidad = "\n"
        precio = "\n"
        dto = "\n"
        importe = "\n"
        for articulo in articulosPresupuesto:
            ref = ref + articulo[0] + "\n"
            articulos = articulos + articulo[1] + "\n"
            cantidad = cantidad + str(articulo[2]) + "\n"
            precio = precio + str(articulo[3]) + "\n"
            if(articulo[4] != 0):
                dto = dto + str(articulo[4]) + "%\n"
            importe = importe + str(articulo[5]) + "\n"

        table.rows[3].cells[0].text = ref
        table.rows[3].cells[1].text = articulos
        table.rows[3].cells[1].merge(table.rows[3].cells[2])
        table.rows[3].cells[1].merge(table.rows[3].cells[3])
        table.rows[3].cells[4].text = cantidad
        table.rows[3].cells[5].text = precio
        table.rows[3].cells[6].text = dto
        table.rows[3].cells[7].text = importe

        table.add_row()
        table.rows[4].cells[1].text = "\nMano de obra"
        table.rows[4].cells[1].merge(table.rows[4].cells[2])
        table.rows[4].cells[1].merge(table.rows[4].cells[3])
        table.rows[4].cells[4].text = "\n" + str(horasManoObra)
        table.rows[4].cells[5].text = "\n" + str(costeMensualHoraBeneficio)
        table.rows[4].cells[7].text = "\n" + str(importeManoObra)

        table.add_row()
        table.rows[5].cells[1].text = "Observaciones:\n\n" + datosPresupuesto.get("observaciones") + "\n"
        table.rows[5].cells[1].merge(table.rows[5].cells[2])
        table.rows[5].cells[1].merge(table.rows[5].cells[3])
        table.rows[5].cells[4].text = "Base imponible\n"
        table.rows[5].cells[4].merge(table.rows[5].cells[5])
        table.rows[5].cells[4].merge(table.rows[5].cells[6])
        table.rows[5].cells[7].text = str(baseImponible) + "\n"
        
        table.add_row()
        table.rows[5].cells[0].merge(table.rows[6].cells[0])
        table.rows[6].cells[1].merge(table.rows[6].cells[2])
        table.rows[6].cells[1].merge(table.rows[6].cells[3])
        table.rows[5].cells[1].merge(table.rows[6].cells[1])
        table.rows[6].cells[4].text = "Importe IVA ( " + str(iva) + " %)\n"
        table.rows[6].cells[4].merge(table.rows[6].cells[5])
        table.rows[6].cells[4].merge(table.rows[6].cells[6])
        table.rows[6].cells[7].text = str(importeIva) + "\n"

        table.add_row()
        table.rows[5].cells[0].merge(table.rows[7].cells[0])
        table.rows[7].cells[1].text = "Cláusulas:\n\n" + datosPresupuesto.get("clausulas") + "\n"
        table.rows[7].cells[1].merge(table.rows[7].cells[2])
        table.rows[7].cells[1].merge(table.rows[7].cells[3])
        table.rows[7].cells[4].text = "Total Presupuesto\n"
        table.rows[7].cells[4].merge(table.rows[7].cells[5])
        table.rows[7].cells[4].merge(table.rows[7].cells[6])
        table.rows[7].cells[7].text = str(totalPresupuesto) + "\n"

        table.add_row()
        table.rows[5].cells[0].merge(table.rows[8].cells[0])
        table.rows[8].cells[1].merge(table.rows[8].cells[2])
        table.rows[8].cells[1].merge(table.rows[8].cells[3])
        table.rows[7].cells[1].merge(table.rows[8].cells[1])
        table.rows[8].cells[4].text = "Firma del cliente\n\n\n"
        table.rows[8].cells[4].merge(table.rows[8].cells[5])
        table.rows[8].cells[4].merge(table.rows[8].cells[6])
        table.rows[8].cells[4].merge(table.rows[8].cells[7])

        presupuesto.save("presupuesto.docx")

        inputFile = "presupuesto.docx"
        outputFile = "presupuesto.pdf"
        file = open(outputFile, "w")
        file.close()

        docx2pdf.convert(inputFile, outputFile)

        #Factura

        datosFactura = {}
        datosFactura["numero"] = input("Introduce el número de la factura: ")
        datosFactura["albaran"] = input("Introduce el número del albarán: ")
        datosFactura["fecha"] = input("Introduce la fecha de la factura: ")
        datosFactura["observaciones"] = input("Introduce las observaciones de la factura: ")
        datosFactura["formaPago"] = input("Introduce la forma de pago: ")

        articulosFactura = []
        for articulo in articulosPresupuesto:
            ref = articulo[0]
            nombre = articulo[1]
            cantidad = articulo[2]
            precio = articulo[3]
            importe = cantidad * precio
            descuentoComercial = str(articulo[4]) + "% " + nombre
            importeDescontado = importe * articulo[4] / 100
            articulosFactura.append([ref, nombre, cantidad, precio, importe, descuentoComercial, importeDescontado])
        
        #Continuar con creación de la plantilla debajo (Factura)

        factura = docx.Document()
        table = factura.add_table(rows=2, cols=6)
        table.style = 'Table Grid'
        if(datosVendedor.get("cif") != None):
            table.rows[0].cells[0].text = "\nDatos del vendedor:\n" + datosVendedor.get("nombre") + "\nCIF: " + datosVendedor.get("cif") + "\n" + datosVendedor.get("calle") + "\nCP " + datosVendedor.get("codigoPostal") + "\n"
        else:
            table.rows[0].cells[0].text = "\nDatos del vendedor:\n" + datosVendedor.get("nombre") + "\nNIF: " + datosVendedor.get("nif") + "\n" + datosVendedor.get("calle") + "\nCP " + datosVendedor.get("codigoPostal") + "\n"
        table.rows[0].cells[0].merge(table.rows[0].cells[1])
        table.rows[0].cells[0].merge(table.rows[0].cells[2])
        table.rows[0].cells[0].merge(table.rows[0].cells[3])
        table.rows[0].cells[0].merge(table.rows[0].cells[4])
        table.rows[0].cells[0].merge(table.rows[0].cells[5])

        table.rows[1].cells[0].text = "\nFactura nº: " + datosFactura.get("numero") + "\nReferencias:\nAlbarán nº: " + datosFactura.get("albaran") + "\nFecha factura: " + datosFactura.get("fecha") + "\n"
        table.rows[1].cells[0].merge(table.rows[1].cells[1])
        table.rows[1].cells[0].merge(table.rows[1].cells[2])
        if(datosCliente.get("cif") != None):
            table.rows[1].cells[3].text = "\nDatos del cliente:\n" + datosCliente.get("nombre") + "\nCIF: " + datosCliente.get("cif") + "\n" + datosCliente.get("calle") + "\nCP " + datosCliente.get("codigoPostal") + "\n"
        else:
            table.rows[1].cells[3].text = "\nDatos del cliente:\n" + datosCliente.get("nombre") + "\nNIF: " + datosCliente.get("nif") + "\n" + datosCliente.get("calle") + "\nCP " + datosCliente.get("codigoPostal") + "\n"
        table.rows[1].cells[3].merge(table.rows[1].cells[4])
        table.rows[1].cells[3].merge(table.rows[1].cells[5])

        table.add_row()
        table.rows[2].cells[0].text = "Ref"
        table.rows[2].cells[1].text = "Artículos"
        table.rows[2].cells[1].merge(table.rows[2].cells[2])
        table.rows[2].cells[3].text = "Cantidad"
        table.rows[2].cells[4].text = "Precio"
        table.rows[2].cells[5].text = "Importe"

        table.add_row()
        ref = "\n"
        articulos = "\n"
        cantidad = "\n"
        precio = "\n"
        importe = "\n"
        dtoComercial = "\nDESCUENTO COMERCIAL\n"
        importeDto = "\n\n"
        for articulo in articulosFactura:
            ref = ref + articulo[0] + "\n"
            articulos = articulos + articulo[1] + "\n"
            cantidad = cantidad + str(articulo[2]) + "\n"
            precio = precio + str(articulo[3]) + "\n"
            importe = importe + str(articulo[4]) + "\n"
            dtoComercial = dtoComercial + articulo[5] + "\n"
            importeDto = importeDto + str(articulo[6]) + "\n"

        table.rows[3].cells[0].text = ref
        table.rows[3].cells[1].text = articulos
        table.rows[3].cells[1].merge(table.rows[3].cells[2])
        table.rows[3].cells[3].text = cantidad
        table.rows[3].cells[4].text = precio
        table.rows[3].cells[5].text = importe

        table.add_row()
        table.rows[4].cells[0].merge(table.rows[4].cells[1])
        table.rows[4].cells[0].merge(table.rows[4].cells[2])
        table.rows[4].cells[0].merge(table.rows[4].cells[3])
        table.rows[4].cells[0].merge(table.rows[4].cells[4])
        table.rows[4].cells[0].text = dtoComercial
        table.rows[4].cells[5].text = importeDto

        table.add_row()
        table.rows[5].cells[0].text = "\nObservaciones:\n" + datosFactura.get("observaciones") + "\n\nForma de pago:" + datosFactura.get("formaPago") + "\n"
        table.rows[5].cells[0].merge(table.rows[5].cells[1])
        table.rows[5].cells[2].text = "Tipo IVA:"
        table.rows[5].cells[3].text = "RE:"
        table.rows[5].cells[4].text = "Base imponible:"
        table.rows[5].cells[5].text = "Importe IVA:"

        table.add_row()
        table.rows[6].cells[0].merge(table.rows[6].cells[1])
        table.rows[5].cells[0].merge(table.rows[6].cells[0])
        table.rows[6].cells[2].text = str(iva) + "%"
        table.rows[6].cells[4].text = str(baseImponible)
        table.rows[6].cells[5].text = str(importeIva)

        table.add_row()
        table.rows[7].cells[0].merge(table.rows[7].cells[1])
        table.rows[5].cells[0].merge(table.rows[7].cells[0])
        table.rows[7].cells[2].text = "\n\nTotal Factura: \n"
        table.rows[7].cells[2].merge(table.rows[7].cells[3])
        table.rows[7].cells[4].text = "\n\n" + str(totalPresupuesto) + "\n"
        table.rows[7].cells[4].merge(table.rows[7].cells[5])

        factura.save("factura.docx")

        inputFile = "factura.docx"
        outputFile = "factura.pdf"
        file = open(outputFile, "w")
        file.close()

        docx2pdf.convert(inputFile, outputFile)